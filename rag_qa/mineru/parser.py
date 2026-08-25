"""MinerU 风格的文档解析适配层。

目标：
- 输入 PDF / HTML / MD / DOCX
- 输出结构化 JSON：blocks[{type, text, page?, bbox?, block_id}]
- 不强制依赖 GPU / MinerU 官方服务，本地 PyMuPDF 兜底
- 兼容 MinerU v0.x 输出 schema，便于未来切换到官方服务

minerU 输出标准（精简）：
{
  "pdf_info": [{"page_idx":0, "page_size":[w,h]}],
  "content": {
    "blocks": [
      {"type":"text","text":"...","page_idx":0,"bbox":[x0,y0,x1,y1]},
      {"type":"table","html":"...","page_idx":1},
      {"type":"image","img_path":"...","caption":"..."}
    ]
  }
}
"""

from __future__ import annotations

import hashlib
import json
import os
import re
import uuid
from dataclasses import dataclass
from typing import Any, Dict, List, Optional

from base import logger


@dataclass
class ParsedBlock:
    block_id: str
    type: str          # text / title / list / table / image
    text: str
    page_idx: int = -1
    bbox: Optional[List[float]] = None
    meta: Dict[str, Any] = None  # type: ignore[assignment]


def _gen_id(text: str, page: int) -> str:
    h = hashlib.md5(f"{page}::{text[:60]}".encode("utf-8")).hexdigest()[:10]
    return f"blk_{page}_{h}"


# ---------- 解析器 ----------
class MinerUStyleParser:
    """MinerU 风格文档解析。

    - 优先调用官方 MinerU（如已部署），否则本地用 PyMuPDF / pdfplumber 兜底。
    - 出来的 block 列表会送进 RAG 的 parent-child 切分器。
    """

    def __init__(self, mineru_endpoint: Optional[str] = None):
        self.mineru_endpoint = mineru_endpoint
        self._pymupdf_ok = self._probe_pymupdf()
        self._pdfplumber_ok = self._probe_pdfplumber()

    def _probe_pymupdf(self) -> bool:
        try:
            import fitz  # noqa: F401
            return True
        except ImportError:
            return False

    def _probe_pdfplumber(self) -> bool:
        try:
            import pdfplumber  # noqa: F401
            return True
        except ImportError:
            return False

    # ---------- 公开 API ----------
    def parse(self, file_path: str) -> Dict[str, Any]:
        """主入口：解析文件，返回 MinerU-style dict。"""
        ext = os.path.splitext(file_path)[-1].lower()
        if ext == ".pdf":
            return self._parse_pdf(file_path)
        if ext in (".md", ".markdown"):
            return self._parse_md(file_path)
        if ext in (".html", ".htm"):
            return self._parse_html(file_path)
        if ext == ".docx":
            return self._parse_docx(file_path)
        if ext == ".txt":
            return self._parse_txt(file_path)
        raise ValueError(f"Unsupported file type: {ext}")

    def parse_bytes(self, content: bytes, filename: str) -> Dict[str, Any]:
        """把上传字节流先存到临时文件再解析。"""
        tmp_dir = os.environ.get("MINERU_TMP_DIR", "/tmp/mineru")
        os.makedirs(tmp_dir, exist_ok=True)
        tmp_path = os.path.join(tmp_dir, f"{uuid.uuid4().hex}_{filename}")
        with open(tmp_path, "wb") as f:
            f.write(content)
        try:
            return self.parse(tmp_path)
        finally:
            try:
                os.remove(tmp_path)
            except OSError:
                pass

    # ---------- 各类解析 ----------
    def _parse_pdf(self, file_path: str) -> Dict[str, Any]:
        """优先 MinerU 官方 API，降级到 PyMuPDF。"""
        if self.mineru_endpoint:
            try:
                return self._parse_via_mineru(file_path)
            except Exception as e:
                logger.warning("MinerU endpoint failed, fallback pymupdf: %s", e)
        if self._pymupdf_ok:
            return self._parse_pdf_pymupdf(file_path)
        if self._pdfplumber_ok:
            return self._parse_pdf_pdfplumber(file_path)
        raise RuntimeError(
            "No PDF parser available. Install pymupdf or set mineru_endpoint."
        )

    def _parse_via_mineru(self, file_path: str) -> Dict[str, Any]:
        import requests  # 临时导入
        with open(file_path, "rb") as f:
            resp = requests.post(
                f"{self.mineru_endpoint.rstrip('/')}/predict",
                files={"file": f}, timeout=60,
            )
        resp.raise_for_status()
        return resp.json()

    def _parse_pdf_pymupdf(self, file_path: str) -> Dict[str, Any]:
        import fitz  # type: ignore
        doc = fitz.open(file_path)
        blocks: List[Dict[str, Any]] = []
        pdf_info: List[Dict[str, Any]] = []
        for page_idx, page in enumerate(doc):
            w, h = page.rect.width, page.rect.height
            pdf_info.append({"page_idx": page_idx, "page_size": [w, h]})
            for b in page.get_text("blocks"):
                x0, y0, x1, y1, text = b[0], b[1], b[2], b[3], b[4]
                cleaned = (text or "").strip()
                if not cleaned:
                    continue
                block_type = "title" if (
                    len(cleaned) < 60 and cleaned.count("\n") == 0
                    and re.match(r"^第?[一二三四五六七八九十0-9]+[章节]", cleaned)
                ) else "text"
                blocks.append({
                    "type": block_type,
                    "text": cleaned,
                    "page_idx": page_idx,
                    "bbox": [x0, y0, x1, y1],
                })
        doc.close()
        return self._wrap(file_path, pdf_info, blocks)

    def _parse_pdf_pdfplumber(self, file_path: str) -> Dict[str, Any]:
        import pdfplumber  # type: ignore
        blocks: List[Dict[str, Any]] = []
        pdf_info: List[Dict[str, Any]] = []
        with pdfplumber.open(file_path) as pdf:
            for page_idx, page in enumerate(pdf.pages):
                w, h = page.width, page.height
                pdf_info.append({"page_idx": page_idx, "page_size": [w, h]})
                # 文字
                text = page.extract_text() or ""
                if text.strip():
                    blocks.append({"type": "text", "text": text.strip(),
                                   "page_idx": page_idx})
                # 表格
                for table in page.extract_tables() or []:
                    rows = ["|".join(str(c or "") for c in r) for r in table]
                    blocks.append({
                        "type": "table",
                        "html": "<table>" + "".join(
                            f"<tr>{''.join(f'<td>{c}</td>' for c in r)}</tr>"
                            for r in table
                        ) + "</table>",
                        "text": "\n".join(rows),
                        "page_idx": page_idx,
                    })
        return self._wrap(file_path, pdf_info, blocks)

    def _parse_md(self, file_path: str) -> Dict[str, Any]:
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()
        lines = text.split("\n")
        blocks: List[Dict[str, Any]] = []
        cur_buf: List[str] = []
        cur_type = "text"
        for line in lines:
            if line.startswith("# "):
                if cur_buf:
                    blocks.append({"type": cur_type,
                                   "text": "\n".join(cur_buf).strip()})
                    cur_buf = []
                blocks.append({"type": "title", "text": line[2:].strip()})
            elif line.startswith("## "):
                if cur_buf:
                    blocks.append({"type": cur_type,
                                   "text": "\n".join(cur_buf).strip()})
                    cur_buf = []
                blocks.append({"type": "title", "text": line[3:].strip()})
            else:
                cur_buf.append(line)
        if cur_buf:
            blocks.append({"type": cur_type, "text": "\n".join(cur_buf).strip()})
        return self._wrap(file_path, [{"page_idx": 0, "page_size": [0, 0]}],
                          blocks)

    def _parse_html(self, file_path: str) -> Dict[str, Any]:
        from bs4 import BeautifulSoup  # type: ignore
        with open(file_path, "r", encoding="utf-8") as f:
            html = f.read()
        soup = BeautifulSoup(html, "html.parser")
        blocks: List[Dict[str, Any]] = []
        for elem in soup.find_all(["h1", "h2", "h3", "p", "li", "table"]):
            text = elem.get_text(" ", strip=True)
            if not text:
                continue
            if elem.name.startswith("h"):
                blocks.append({"type": "title", "text": text})
            elif elem.name == "li":
                blocks.append({"type": "list", "text": text})
            elif elem.name == "table":
                blocks.append({"type": "table", "text": text,
                               "html": str(elem)})
            else:
                blocks.append({"type": "text", "text": text})
        return self._wrap(file_path, [{"page_idx": 0, "page_size": [0, 0]}],
                          blocks)

    def _parse_docx(self, file_path: str) -> Dict[str, Any]:
        try:
            from docx import Document  # type: ignore
        except ImportError as e:
            raise RuntimeError("docx not installed") from e
        doc = Document(file_path)
        blocks: List[Dict[str, Any]] = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            style = (para.style.name or "").lower() if para.style else ""
            btype = "title" if "heading" in style else "text"
            blocks.append({"type": btype, "text": text})
        for table in doc.tables:
            rows = []
            for r in table.rows:
                rows.append("|".join(c.text for c in r.cells))
            blocks.append({
                "type": "table",
                "text": "\n".join(rows),
                "html": "",
            })
        return self._wrap(file_path, [{"page_idx": 0, "page_size": [0, 0]}],
                          blocks)

    def _parse_txt(self, file_path: str) -> Dict[str, Any]:
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()
        blocks = [{"type": "text", "text": text.strip()}]
        return self._wrap(file_path, [{"page_idx": 0, "page_size": [0, 0]}],
                          blocks)

    # ---------- 包装 ----------
    def _wrap(self, file_path: str,
              pdf_info: List[Dict[str, Any]],
              blocks: List[Dict[str, Any]]) -> Dict[str, Any]:
        # 给每个 block 加 block_id
        for b in blocks:
            page = b.get("page_idx", 0)
            txt = b.get("text") or b.get("html") or ""
            b["block_id"] = _gen_id(txt, page)
        return {
            "pdf_info": pdf_info,
            "content": {"blocks": blocks},
            "meta": {
                "source_file": os.path.basename(file_path),
                "parser": "mineru_style",
                "block_count": len(blocks),
            },
        }

    def to_plain_documents(self, parsed: Dict[str, Any]
                           ) -> List[Dict[str, Any]]:
        """把 MinerU-style 输出展平成 {'chunk_id','text','meta'} 列表，
        供 RAG 切分器使用。
        """
        out: List[Dict[str, Any]] = []
        for b in parsed.get("content", {}).get("blocks", []):
            text = b.get("text") or ""
            if not text:
                continue
            out.append({
                "chunk_id": b["block_id"],
                "text": text,
                "meta": {
                    "type": b.get("type", "text"),
                    "page_idx": b.get("page_idx", 0),
                    "source": parsed.get("meta", {}).get("source_file", ""),
                },
            })
        return out

    # ---------- 引用查找 ----------
    def fetch_chunk(self, chunk_id: str) -> str:
        """Agent 工具 doc_lookup 调用：根据 chunk_id 拿到原文。
        这要求解析时把 blocks 缓存到磁盘；这里做一个最简实现。
        """
        cache_path = os.environ.get("MINERU_CACHE", "/tmp/mineru_cache")
        if not os.path.isdir(cache_path):
            return ""
        for fname in os.listdir(cache_path):
            fp = os.path.join(cache_path, fname)
            try:
                with open(fp, "r", encoding="utf-8") as f:
                    data = json.load(f)
                for b in data.get("content", {}).get("blocks", []):
                    if b.get("block_id") == chunk_id:
                        return b.get("text", "")
            except Exception:
                continue
        return ""
